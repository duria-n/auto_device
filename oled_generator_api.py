"""
oled_generator_api.py
=====================
OLED器件例专利文本生成器 v2 — 统一外部调用接口

支持：
  - Python 直接 import 调用
  - 命令行 CLI 调用
  - HTTP 客户端调用（Flask 服务需已启动）

模型支持：OpenAI / Anthropic / Qwen / Deepseek

快速示例：
  >>> from oled_generator_api import generate_from_csv, OLEDClient
  >>> # 纯模板，无需 API Key
  >>> generate_from_csv("data/devices.csv", output_format="docx_file", output_dir="out/")
  >>> # LLM 润色
  >>> generate_from_csv("data/devices.csv", mode="template_polish",
  ...     llm={"provider":"qwen","model":"qwen-max","api_key":"sk-xxx"},
  ...     output_format="docx_file", output_dir="out/")
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

import pandas as pd

# ── 项目根目录加入 sys.path ───────────────────────────────────────────
_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))

from core.parser import parse_row, parse_dataframe, parse_csv as _parse_csv
from core.constants import DEFAULT_SECTIONS
from llm.base import LLMConfig
from llm.service import generate as _generate

# ── 正则 ──────────────────────────────────────────────────────────────
_PH_RE    = re.compile(r"【待补充:[^】]+】")
_TITLE_RE = re.compile(r"^[①②③④⑤⑥]|^对比例\s*\d+\s*$|^实施例\s*\d+\s*$")


# ══════════════════════════════════════════════════════════════════════
# 内部工具
# ══════════════════════════════════════════════════════════════════════

def _make_llm_cfg(llm: dict | None) -> LLMConfig | None:
    """将 dict 形式的 LLM 配置转为 LLMConfig。llm=None 时走纯模板模式。"""
    if not llm:
        return None
    provider = llm.get("provider", "openai").lower()

    cfg = LLMConfig(
        provider    = provider,
        model       = llm.get("model", ""),
        api_key     = llm.get("api_key", ""),
        max_tokens  = int(llm.get("max_tokens", 2048)),
        temperature = float(llm.get("temperature", 0.3)),
    )

    if provider == "huggingface":
        cfg.hf_cache_dir      = llm.get("hf_cache_dir",      "")
        cfg.hf_device         = llm.get("hf_device",         "auto")
        cfg.hf_torch_dtype    = llm.get("hf_torch_dtype",    "auto")
        cfg.hf_load_in_4bit   = bool(llm.get("hf_load_in_4bit",  False))
        cfg.hf_load_in_8bit   = bool(llm.get("hf_load_in_8bit",  False))
        cfg.hf_token          = llm.get("hf_token",          "")
        cfg.hf_max_new_tokens = int(llm.get("hf_max_new_tokens", 2048))
        if not cfg.model:
            return None   # HF 模式必须指定 model_id
    else:
        if not cfg.api_key:
            return None   # 非HF模式必须有 api_key

    return cfg


def _write_txt(text: str, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    print(f"  ✓ TXT  → {path}")
    return path


def _write_docx(text: str, path: Path) -> Path:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_LINE_SPACING
    except ImportError:
        raise ImportError("pip install python-docx")

    ORANGE = RGBColor(0xFF, 0x8C, 0x00)
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.54)
        sec.left_margin = sec.right_margin  = Cm(3.18)
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _r(para, t, bold=False, color=None):
        run = para.add_run(t)
        run.font.name = "宋体"; run.font.size = Pt(12); run.font.bold = bold
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if color: run.font.color.rgb = color

    for line in text.split("\n"):
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = Pt(18)
        if not line.strip(): continue
        bold = bool(_TITLE_RE.match(line.strip()))
        for i, seg in enumerate(_PH_RE.split(line)):
            if seg: _r(para, seg, bold=bold)
            phs = _PH_RE.findall(line)
            if i < len(phs): _r(para, phs[i], color=ORANGE)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"  ✓ DOCX → {path}")
    return path


def _dispatch_write(text: str, fmt: str, path: Path | None) -> Path | None:
    if fmt == "txt_file":
        if not path: raise ValueError("output_path 不能为空")
        return _write_txt(text, path)
    if fmt == "docx_file":
        if not path: raise ValueError("output_path 不能为空")
        return _write_docx(text, path)
    return None


# ══════════════════════════════════════════════════════════════════════
# 公开 API
# ══════════════════════════════════════════════════════════════════════

def generate_from_dict(
    data: dict[str, Any],
    sections: list[str] | None = None,
    mode: str = "template",
    llm: dict | None = None,
    output_format: str = "text",
    output_path: str | Path | None = None,
) -> str:
    """
    从 dict 生成单条器件例文本。

    Parameters
    ----------
    data : dict
        器件配方，需含 "materials" 键（list[dict]）。
    sections : list[str], optional
        要生成的段落，默认全部。
    mode : str
        "template"        — 纯模板，无需 llm 参数
        "template_polish" — 模板草稿 + LLM 润色
        "llm_direct"      — LLM 直接撰写
    llm : dict, optional
        LLM 配置：{"provider","model","api_key","max_tokens","temperature"}
        mode 为 "template" 时忽略。
    output_format : str
        "text" | "txt_file" | "docx_file"
    output_path : str | Path, optional
        输出文件路径（output_format 为文件格式时必填）。

    Returns
    -------
    str  生成的纯文本。

    Examples
    --------
    >>> text = generate_from_dict(data)
    >>> text = generate_from_dict(data, mode="template_polish",
    ...     llm={"provider":"openai","model":"gpt-4o","api_key":"sk-xxx"})
    >>> generate_from_dict(data, output_format="docx_file", output_path="out/实施例1.docx")
    """
    recipe   = parse_row(data)
    llm_cfg  = _make_llm_cfg(llm)
    text     = _generate(recipe, sections or DEFAULT_SECTIONS, llm_cfg, mode)
    _dispatch_write(text, output_format, Path(output_path) if output_path else None)
    return text


def generate_from_dataframe(
    df: pd.DataFrame,
    sections: list[str] | None = None,
    mode: str = "template",
    llm: dict | None = None,
    output_format: str = "text",
    output_dir: str | Path | None = None,
    filename_col: str | None = None,
) -> list[dict[str, Any]]:
    """
    从 DataFrame 批量生成器件例文本。

    Parameters
    ----------
    df : pd.DataFrame
        每行一条器件配方，列名遵循 CSV 模板格式。
    sections / mode / llm / output_format
        同 generate_from_dict。
    output_dir : str | Path, optional
        批量写文件的输出目录。
    filename_col : str, optional
        用于命名输出文件的列名，默认使用 device_no。

    Returns
    -------
    list[dict]
        每项含：{"index", "device_no", "text", "file"}

    Examples
    --------
    >>> df = pd.read_csv("data/devices.csv")
    >>> results = generate_from_dataframe(df, output_format="docx_file", output_dir="out/")
    >>> results = generate_from_dataframe(df, mode="template_polish",
    ...     llm={"provider":"deepseek","model":"deepseek-chat","api_key":"sk-xxx"},
    ...     output_format="docx_file", output_dir="out/")
    """
    secs    = sections or DEFAULT_SECTIONS
    llm_cfg = _make_llm_cfg(llm)
    out_dir = Path(output_dir) if output_dir else None
    if out_dir and output_format != "text":
        out_dir.mkdir(parents=True, exist_ok=True)

    results: list[dict[str, Any]] = []
    recipes = parse_dataframe(df)

    for i, recipe in enumerate(recipes):
        text     = _generate(recipe, secs, llm_cfg, mode)
        fname    = str(df.iloc[i].get(filename_col, recipe.device_no)) if filename_col \
                   else recipe.device_no
        out_file = None
        if out_dir and output_format != "text":
            ext      = "txt" if output_format == "txt_file" else "docx"
            out_file = out_dir / f"{fname}.{ext}"
            _dispatch_write(text, output_format, out_file)
        results.append({"index": i+1, "device_no": recipe.device_no, "text": text, "file": out_file})

    return results


def generate_from_csv(
    csv_path: str | Path,
    sections: list[str] | None = None,
    mode: str = "template",
    llm: dict | None = None,
    output_format: str = "text",
    output_dir: str | Path | None = None,
    encoding: str = "utf-8-sig",
    filename_col: str | None = None,
) -> list[dict[str, Any]]:
    """
    从 CSV / Excel 文件批量生成器件例文本。

    Parameters
    ----------
    csv_path : str | Path   .csv / .xlsx / .xls 文件路径。
    sections / mode / llm / output_format / output_dir / filename_col
        同 generate_from_dataframe。
    encoding : str          CSV 编码，默认 utf-8-sig。

    Examples
    --------
    >>> # 纯模板，批量 DOCX
    >>> generate_from_csv("data/devices.csv",
    ...     output_format="docx_file", output_dir="out/")

    >>> # Qwen 润色，批量 DOCX
    >>> generate_from_csv("data/devices.csv",
    ...     mode="template_polish",
    ...     llm={"provider":"qwen","model":"qwen-max","api_key":"sk-xxx"},
    ...     output_format="docx_file", output_dir="out/")

    >>> # 只取文本，不写文件
    >>> results = generate_from_csv("data/devices.csv")
    >>> texts = [r["text"] for r in results]
    """
    p  = Path(csv_path)
    df = pd.read_excel(p) if p.suffix.lower() in (".xlsx", ".xls") \
         else pd.read_csv(p, encoding=encoding)
    return generate_from_dataframe(df, sections=sections, mode=mode, llm=llm,
                                   output_format=output_format, output_dir=output_dir,
                                   filename_col=filename_col)


# ── 快捷方式 ──────────────────────────────────────────────────────────
def generate_single_to_docx(data, path, sections=None, mode="template", llm=None):
    """dict → DOCX 快捷方式。"""
    return generate_from_dict(data, sections=sections, mode=mode, llm=llm,
                              output_format="docx_file", output_path=path)

def generate_single_to_txt(data, path, sections=None, mode="template", llm=None):
    """dict → TXT 快捷方式。"""
    return generate_from_dict(data, sections=sections, mode=mode, llm=llm,
                              output_format="txt_file", output_path=path)

def generate_batch_to_docx(source, output_dir, sections=None, mode="template", llm=None):
    """CSV路径或DataFrame → 批量 DOCX 快捷方式。"""
    if isinstance(source, pd.DataFrame):
        return generate_from_dataframe(source, sections=sections, mode=mode, llm=llm,
                                       output_format="docx_file", output_dir=output_dir)
    return generate_from_csv(source, sections=sections, mode=mode, llm=llm,
                             output_format="docx_file", output_dir=output_dir)

def generate_batch_to_txt(source, output_dir, sections=None, mode="template", llm=None):
    """CSV路径或DataFrame → 批量 TXT 快捷方式。"""
    if isinstance(source, pd.DataFrame):
        return generate_from_dataframe(source, sections=sections, mode=mode, llm=llm,
                                       output_format="txt_file", output_dir=output_dir)
    return generate_from_csv(source, sections=sections, mode=mode, llm=llm,
                             output_format="txt_file", output_dir=output_dir)


# ══════════════════════════════════════════════════════════════════════
# HTTP 客户端
# ══════════════════════════════════════════════════════════════════════

class OLEDClient:
    """
    HTTP 客户端，对接本地 Flask 服务（python app.py 启动后使用）。

    Examples
    --------
    >>> client = OLEDClient("http://localhost:5000")
    >>> text = client.generate(data)
    >>> client.save_docx(text, "out/实施例1.docx")
    >>> results = client.batch("data/devices.csv",
    ...     mode="template_polish",
    ...     llm={"provider":"anthropic","model":"claude-sonnet-4-5","api_key":"sk-ant-xxx"})
    >>> client.save_batch_docx(results, "out/docx/")
    """

    def __init__(self, base_url: str = "http://localhost:5000", timeout: int = 120):
        try:
            import requests as _req; self._req = _req
        except ImportError:
            raise ImportError("pip install requests")
        self.base_url = base_url.rstrip("/")
        self.timeout  = timeout

    def generate(
        self,
        data: dict[str, Any],
        sections: list[str] | None = None,
        mode: str = "template",
        llm: dict | None = None,
    ) -> str:
        """POST /api/generate — 单条生成。"""
        payload = {
            "data": data,
            "sections": sections or DEFAULT_SECTIONS,
            "mode": mode,
        }
        if llm and mode != "template":
            payload["llm"] = llm
        resp = self._req.post(f"{self.base_url}/api/generate", json=payload, timeout=self.timeout)
        resp.raise_for_status()
        r = resp.json()
        if not r.get("success"): raise RuntimeError(r.get("error"))
        return r["text"]

    def batch(
        self,
        file_path: str | Path,
        sections: list[str] | None = None,
        mode: str = "template",
        llm: dict | None = None,
    ) -> list[dict[str, Any]]:
        """POST /api/batch — CSV 批量生成。"""
        form: dict[str, Any] = {
            "sections": ",".join(sections or DEFAULT_SECTIONS),
            "mode": mode,
        }
        if llm and mode != "template":
            form.update({
                "api_key":     llm.get("api_key", ""),
                "provider":    llm.get("provider", "openai"),
                "model":       llm.get("model", "gpt-4o"),
                "max_tokens":  str(llm.get("max_tokens", 4096)),
                "temperature": str(llm.get("temperature", 0.3)),
            })
        with open(file_path, "rb") as f:
            resp = self._req.post(f"{self.base_url}/api/batch",
                                  files={"file": f}, data=form,
                                  timeout=max(self.timeout, 300))
        resp.raise_for_status()
        r = resp.json()
        if not r.get("success"): raise RuntimeError(r.get("error"))
        return r["results"]

    def save_txt(self, text: str, path: str | Path) -> Path:
        return _write_txt(text, Path(path))

    def save_docx(self, text: str, path: str | Path) -> Path:
        return _write_docx(text, Path(path))

    def save_batch_txt(self, results: list[dict], output_dir: str | Path) -> list[Path]:
        out = Path(output_dir); out.mkdir(parents=True, exist_ok=True)
        return [_write_txt(r["text"], out/f"{r.get('device_no',r['index'])}.txt") for r in results]

    def save_batch_docx(self, results: list[dict], output_dir: str | Path) -> list[Path]:
        out = Path(output_dir); out.mkdir(parents=True, exist_ok=True)
        return [_write_docx(r["text"], out/f"{r.get('device_no',r['index'])}.docx") for r in results]


# ══════════════════════════════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════════════════════════════

def _cli():
    p = argparse.ArgumentParser(
        prog="oled_generator_api",
        description="OLED器件例专利文本生成器 CLI v2",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  # 纯模板批量 DOCX（无需 API Key）
  python oled_generator_api.py batch -i data/devices.csv -o out/ -f docx_file

  # Qwen 润色批量 DOCX
  python oled_generator_api.py batch -i data/devices.csv -o out/ -f docx_file \\
    --mode template_polish --provider qwen --model qwen-max --api-key sk-xxx

  # Deepseek 直写，只含结构+性能段
  python oled_generator_api.py batch -i data/devices.xlsx -o out/ -f docx_file \\
    --mode llm_direct --provider deepseek --model deepseek-chat --api-key sk-xxx \\
    --sections structure performance

  # 单条 JSON → DOCX
  python oled_generator_api.py single -i data/device1.json -o out/实施例1.docx -f docx_file
""")

    sub = p.add_subparsers(dest="cmd", required=True)

    def _common(sp):
        sp.add_argument("--mode", default="template",
                        choices=["template","template_polish","llm_direct"])
        sp.add_argument("--provider", default="openai",
                        choices=["openai","anthropic","qwen","deepseek"])
        sp.add_argument("--model",   default="")
        sp.add_argument("--api-key", default="", dest="api_key")
        sp.add_argument("--max-tokens",  type=int,   default=4096, dest="max_tokens")
        sp.add_argument("--temperature", type=float, default=0.3)
        sp.add_argument("--sections", "-s", nargs="*", default=DEFAULT_SECTIONS,
                        choices=DEFAULT_SECTIONS, metavar="SEC")

    # single
    ps = sub.add_parser("single", help="从 JSON 生成单条")
    ps.add_argument("--input",  "-i", required=True)
    ps.add_argument("--output", "-o", default=None)
    ps.add_argument("--format", "-f", default="text",
                    choices=["text","txt_file","docx_file"])
    _common(ps)

    # batch
    pb = sub.add_parser("batch", help="从 CSV/Excel 批量生成")
    pb.add_argument("--input",       "-i", required=True)
    pb.add_argument("--output-dir",  "-o", default="output", dest="output_dir")
    pb.add_argument("--format",      "-f", default="txt_file",
                    choices=["text","txt_file","docx_file"])
    pb.add_argument("--encoding", default="utf-8-sig")
    pb.add_argument("--filename-col", default=None, dest="filename_col")
    _common(pb)

    args = p.parse_args()

    llm = None
    if args.mode != "template" and args.api_key:
        llm = {
            "provider":    args.provider,
            "model":       args.model,
            "api_key":     args.api_key,
            "max_tokens":  args.max_tokens,
            "temperature": args.temperature,
        }

    if args.cmd == "single":
        with open(args.input, encoding="utf-8") as f:
            data = json.load(f)
        if args.format in ("txt_file","docx_file") and not args.output:
            p.error("--format 为文件格式时必须提供 --output")
        text = generate_from_dict(data, sections=args.sections, mode=args.mode, llm=llm,
                                  output_format=args.format,
                                  output_path=args.output)
        if args.format == "text": print(text)

    elif args.cmd == "batch":
        results = generate_from_csv(
            args.input, sections=args.sections, mode=args.mode, llm=llm,
            output_format=args.format, output_dir=args.output_dir,
            encoding=args.encoding, filename_col=args.filename_col,
        )
        print(f"\n✓ 共生成 {len(results)} 条")
        for r in results:
            suffix = f" → {r['file']}" if r.get("file") else ""
            print(f"  [{r['index']:>3}] {r['device_no']}{suffix}")


if __name__ == "__main__":
    _cli()
