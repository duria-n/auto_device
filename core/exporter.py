"""core/exporter.py — 将生成的文本写入 TXT / DOCX 文件"""
from __future__ import annotations
import re
from pathlib import Path

_PH_RE    = re.compile(r"【待补充:[^】]+】")
_TITLE_RE = re.compile(r"^[①②③④⑤⑥]|^对比例\s*\d+\s*$|^实施例\s*\d+\s*$")


def write_txt(text: str, path: str | Path) -> Path:
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(text, encoding="utf-8")
    return p


def write_docx(text: str, path: str | Path) -> Path:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_LINE_SPACING
    except ImportError:
        raise ImportError("pip install python-docx")

    ORANGE = RGBColor(0xFF, 0x8C, 0x00)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.54)
        sec.left_margin = sec.right_margin = Cm(3.18)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _run(para, text_: str, bold: bool = False, color: RGBColor | None = None):
        run = para.add_run(text_)
        run.font.name = "宋体"
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if color:
            run.font.color.rgb = color

    def _spacing(para):
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = Pt(18)

    for line in text.split("\n"):
        para = doc.add_paragraph()
        _spacing(para)
        if not line.strip():
            continue
        is_title = bool(_TITLE_RE.match(line.strip()))
        segs = _PH_RE.split(line)
        phs  = _PH_RE.findall(line)
        for i, seg in enumerate(segs):
            if seg:
                _run(para, seg, bold=is_title)
            if i < len(phs):
                _run(para, phs[i], color=ORANGE)

    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(p))
    return p
